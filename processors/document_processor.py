import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
import pdfplumber
from docx import Document
import logging

class DocumentProcessor:
    """Processes PDF and DOCX documents for RAG system"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def process_document(self, file_path: str, filename: str) -> List[Dict]:
        """Process a document and return chunks with metadata"""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path, filename)
        elif file_extension == '.docx':
            return self._process_docx(file_path, filename)
        else:
            raise ValueError(f"Unsupported document format: {file_extension}")
    
    def _process_pdf(self, file_path: str, filename: str) -> List[Dict]:
        """Extract text from PDF and create chunks"""
        chunks = []
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        # Split into paragraphs
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        
                        for para_num, paragraph in enumerate(paragraphs):
                            if len(paragraph) > 50:  # Only include substantial paragraphs
                                chunk = {
                                    'type': 'text',
                                    'content': paragraph,
                                    'source_file': filename,
                                    'page': page_num,
                                    'paragraph': para_num + 1,
                                    'chunk_id': f"{filename}_page_{page_num}_para_{para_num + 1}"
                                }
                                chunks.append(chunk)
        
        except Exception as e:
            self.logger.error(f"Error processing PDF {filename}: {str(e)}")
            # Fallback to PyPDF2 if pdfplumber fails
            try:
                chunks = self._process_pdf_fallback(file_path, filename)
            except Exception as e2:
                self.logger.error(f"PDF fallback also failed for {filename}: {str(e2)}")
                raise e2
        
        return chunks
    
    def _process_pdf_fallback(self, file_path: str, filename: str) -> List[Dict]:
        """Fallback PDF processing using PyPDF2"""
        chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    # Clean up text
                    text = text.replace('\n', ' ').strip()
                    
                    # Split into sentences for better chunking
                    sentences = text.split('.')
                    current_chunk = ""
                    chunk_num = 1
                    
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if not sentence:
                            continue
                        
                        # If adding this sentence would make chunk too long, start new chunk
                        if len(current_chunk + sentence) > 1000:
                            if current_chunk:
                                chunk = {
                                    'type': 'text',
                                    'content': current_chunk.strip(),
                                    'source_file': filename,
                                    'page': page_num,
                                    'chunk_number': chunk_num,
                                    'chunk_id': f"{filename}_page_{page_num}_chunk_{chunk_num}"
                                }
                                chunks.append(chunk)
                                chunk_num += 1
                                current_chunk = ""
                        
                        current_chunk += sentence + ". "
                    
                    # Add final chunk if exists
                    if current_chunk.strip():
                        chunk = {
                            'type': 'text',
                            'content': current_chunk.strip(),
                            'source_file': filename,
                            'page': page_num,
                            'chunk_number': chunk_num,
                            'chunk_id': f"{filename}_page_{page_num}_chunk_{chunk_num}"
                        }
                        chunks.append(chunk)
        
        return chunks
    
    def _process_docx(self, file_path: str, filename: str) -> List[Dict]:
        """Extract text from DOCX and create chunks"""
        chunks = []
        
        try:
            doc = Document(file_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text.strip()
                if text and len(text) > 50:  # Only include substantial paragraphs
                    chunk = {
                        'type': 'text',
                        'content': text,
                        'source_file': filename,
                        'paragraph': para_num,
                        'chunk_id': f"{filename}_para_{para_num}"
                    }
                    chunks.append(chunk)
            
            # Also process tables if they exist
            for table_num, table in enumerate(doc.tables, 1):
                table_text = self._extract_table_text(table)
                if table_text:
                    chunk = {
                        'type': 'text',
                        'content': f"Table {table_num}: {table_text}",
                        'source_file': filename,
                        'table': table_num,
                        'chunk_id': f"{filename}_table_{table_num}"
                    }
                    chunks.append(chunk)
        
        except Exception as e:
            self.logger.error(f"Error processing DOCX {filename}: {str(e)}")
            raise e
        
        return chunks
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
